from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
HTML_PATH = ROOT / "docs" / "Report_updated_source.html"
DOCX_PATH = ROOT / "docs" / "report.docx"


def normalize_text(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def extract_inline(paragraph, node: Tag) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                paragraph.add_run(text)
            continue

        if not isinstance(child, Tag):
            continue

        name = child.name.lower()
        if name == "br":
            paragraph.add_run("\n")
            continue

        run = paragraph.add_run(normalize_text(child.get_text(" ", strip=True)))
        if name in {"strong", "b"}:
            run.bold = True
        elif name in {"em", "i"}:
            run.italic = True
        elif name == "code":
            run.font.name = "Courier New"


def resolve_image_path(src: str) -> Path | None:
    raw = (ROOT / src).resolve() if not Path(src).is_absolute() else Path(src)
    if raw.suffix.lower() == ".svg":
        png_candidate = raw.parent / "png" / f"{raw.stem}.png"
        if png_candidate.exists():
            return png_candidate
    if raw.exists():
        return raw
    return None


def add_heading(doc: Document, text: str, level: int) -> None:
    if not text:
        return
    level = max(1, min(4, level))
    doc.add_heading(text, level=level)


def add_table(doc: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return

    max_cols = max(len(r.find_all(["th", "td"])) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for c_idx, cell in enumerate(cells):
            text = normalize_text(cell.get_text(" ", strip=True))
            paragraph = table.cell(r_idx, c_idx).paragraphs[0]
            paragraph.text = text
            if cell.name.lower() == "th":
                for run in paragraph.runs:
                    run.bold = True


def process_node(doc: Document, node: Tag) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            text = normalize_text(str(child))
            if text:
                doc.add_paragraph(text)
            continue

        if not isinstance(child, Tag):
            continue

        name = child.name.lower()

        if name == "section":
            process_node(doc, child)
            doc.add_page_break()
            continue

        if name == "h1":
            add_heading(doc, normalize_text(child.get_text(" ", strip=True)), 1)
            continue
        if name == "h2":
            add_heading(doc, normalize_text(child.get_text(" ", strip=True)), 1)
            continue
        if name == "h3":
            add_heading(doc, normalize_text(child.get_text(" ", strip=True)), 2)
            continue
        if name == "h4":
            add_heading(doc, normalize_text(child.get_text(" ", strip=True)), 3)
            continue

        if name == "p":
            paragraph = doc.add_paragraph()
            extract_inline(paragraph, child)
            continue

        if name in {"ul", "ol"}:
            style = "List Bullet" if name == "ul" else "List Number"
            for li in child.find_all("li", recursive=False):
                paragraph = doc.add_paragraph(style=style)
                extract_inline(paragraph, li)
            continue

        if name == "table":
            add_table(doc, child)
            continue

        if name == "figure":
            img = child.find("img")
            caption = child.find("figcaption")
            if img and img.get("src"):
                image_path = resolve_image_path(img["src"])
                if image_path:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(str(image_path), width=Inches(6.2))
            if caption:
                c = doc.add_paragraph(normalize_text(caption.get_text(" ", strip=True)))
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in c.runs:
                    run.italic = True
            continue

        if name == "div" and "caption-only" in child.get("class", []):
            caption_text = normalize_text(child.get_text(" ", strip=True))
            if caption_text:
                c = doc.add_paragraph(caption_text)
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in c.runs:
                    run.italic = True
            continue

        if name == "pre":
            pre_text = child.get_text("\n", strip=False)
            p = doc.add_paragraph(pre_text)
            for run in p.runs:
                run.font.name = "Courier New"
            continue

        process_node(doc, child)


def build_docx() -> None:
    html = HTML_PATH.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    body = soup.find("body")
    if body is None:
        raise RuntimeError("No body found in report HTML")

    process_node(doc, body)

    if doc.paragraphs and not doc.paragraphs[-1].text.strip():
        pass

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
